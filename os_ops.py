import os
import platform
import winshell
import pyautogui
import screen_brightness_control as sbc
from docx import Document
import pyperclip
import time
import base64
from io import BytesIO


def empty_recycle_bin():
    """Empties the Windows Recycle Bin."""
    if platform.system() == "Windows":
        try:
            winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=True)
            return "Recycle bin emptied successfully, sir."
        except:
            return "The recycle bin is already empty."
    return "Not supported on this OS."


def control_system_hardware(action, level=None):
    """Controls system volume, brightness, or opens settings."""
    if action == "volume_up":
        pyautogui.press("volumeup", presses=10)
        return "Volume increased."
    elif action == "volume_down":
        pyautogui.press("volumedown", presses=10)
        return "Volume decreased."
    elif action == "mute":
        pyautogui.press("volumemute")
        return "System muted."
    elif action == "brightness" and level is not None:
        sbc.set_brightness(level)
        return f"Brightness set to {level} percent."
    elif action == "open_settings":
        os.system("start ms-settings:")
        return "Opening Windows Settings."
    return "Hardware command not recognized."


def create_word_doc(topic, filename, groq_client):
    """Generates an academic assignment and saves it as a Word document."""
    try:
        response = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional academic writer. Write a comprehensive, well-structured assignment. IMPORTANT: Output STRICTLY in plain text with natural paragraphs. DO NOT use markdown, asterisks, hashes, or bullet points.",
                },
                {
                    "role": "user",
                    "content": f"Write an assignment on the topic: {topic}",
                },
            ],
            model="llama-3.1-8b-instant",
        )
        content = response.choices[0].message.content
        clean_content = content.replace("*", "").replace("#", "")

        doc = Document()
        doc.add_heading(topic.title(), 0)
        doc.add_paragraph(clean_content)

        safe_filename = filename.replace(" ", "_") + ".docx"
        doc.save(safe_filename)
        os.startfile(safe_filename)
        return f"Assignment on {topic} has been structured and saved flawlessly."
    except Exception as e:
        return f"Error creating document: {e}"


def open_application(app_name):
    """Smart application launcher."""
    app_name = app_name.lower()
    try:
        if "whatsapp" in app_name:
            os.system("start whatsapp:")
        elif "chrome" in app_name:
            os.system("start chrome")
        else:
            os.system(f"start {app_name}")
        return f"Opening {app_name}, sir."
    except Exception as e:
        return f"Could not open {app_name}."


def read_file_content(file_name):
    """Searches for and reads the contents of a specified file."""
    search_paths = [
        os.getcwd(),
        os.path.expanduser("~\\Desktop"),
        os.path.expanduser("~\\Documents"),
    ]
    target_path = None

    for path in search_paths:
        for root, dirs, files in os.walk(path):
            for file in files:
                if file_name.lower() in file.lower():
                    target_path = os.path.join(root, file)
                    break
            if target_path:
                break
        if target_path:
            break

    if not target_path:
        return f"Sorry sir, I could not find the file named {file_name}."

    try:
        if target_path.endswith(".docx"):
            doc = Document(target_path)
            content = " ".join(
                [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            )
        elif target_path.endswith(".txt"):
            with open(target_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            return "Sir, I can only read text or Word documents."

        if len(content) > 1000:
            content = (
                content[:1000]
                + "... Sir, the file is quite long, I have read the main summary."
            )

        return f"Reading the contents of {file_name}: {content}"
    except Exception as e:
        return f"Error reading file: {e}"


def read_active_window():
    """Copies and reads text from the currently active screen."""
    try:
        pyautogui.hotkey("ctrl", "a")
        time.sleep(0.5)
        pyautogui.hotkey("ctrl", "c")
        time.sleep(0.5)
        pyautogui.press("right")

        copied_text = pyperclip.paste()
        if not copied_text or copied_text.isspace():
            return "Sir, I couldn't find any readable text on the active screen."

        if len(copied_text) > 1000:
            copied_text = copied_text[:1000] + "... [Text truncated]"
        return f"Here is the text from the active screen: {copied_text}"
    except Exception as e:
        return f"Failed to read screen: {e}"


def capture_screen_for_vision():
    """Captures the screen and encodes it in base64 for Vision AI."""
    try:
        screenshot = pyautogui.screenshot()
        buffered = BytesIO()
        screenshot.save(buffered, format="JPEG", quality=70)
        return base64.b64encode(buffered.getvalue()).decode("utf-8")
    except Exception:
        return None


def delete_system_file(file_name):
    """Searches for and permanently deletes a file."""
    search_paths = [
        os.getcwd(),
        os.path.expanduser("~\\Desktop"),
        os.path.expanduser("~\\Documents"),
    ]
    target_path = None

    for path in search_paths:
        for root, dirs, files in os.walk(path):
            for file in files:
                if file_name.lower() in file.lower():
                    target_path = os.path.join(root, file)
                    break
            if target_path:
                break
        if target_path:
            break

    if target_path:
        try:
            os.remove(target_path)
            return f"The file {file_name} has been permanently deleted, sir."
        except Exception as e:
            return f"Error deleting file: {e}"
    return f"I could not find a file named {file_name} to delete."


def edit_word_doc(file_name, new_instruction, groq_client):
    """Appends new AI-generated content to an existing Word document."""
    search_paths = [
        os.getcwd(),
        os.path.expanduser("~\\Desktop"),
        os.path.expanduser("~\\Documents"),
    ]
    target_path = None

    for path in search_paths:
        for root, dirs, files in os.walk(path):
            for file in files:
                if file_name.lower() in file.lower():
                    target_path = os.path.join(root, file)
                    break
            if target_path:
                break
        if target_path:
            break

    if not target_path or not target_path.endswith(".docx"):
        return f"Could not find a valid Word document named {file_name}."

    try:
        response = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional writer. Write a short, structured addition to an ongoing assignment based on the user's prompt. NO markdown formatting.",
                },
                {"role": "user", "content": new_instruction},
            ],
            model="llama-3.1-8b-instant",
        )
        new_content = (
            response.choices[0].message.content.replace("*", "").replace("#", "")
        )

        doc = Document(target_path)
        doc.add_paragraph("\n[Added Update]")
        doc.add_paragraph(new_content)
        doc.save(target_path)

        os.startfile(target_path)
        return f"I have generated and appended the new content to {file_name}."
    except Exception as e:
        return f"Error editing document: {e}"
